"""Add email + LinkedIn hyperlinks to the CV contact line, then export PDF."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parent))
from cv_hyperlinks import DEFAULT_EMAIL, ensure_paragraph_centered, fill_contact_line  # noqa: E402

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "assets" / "files" / "Resume (William Amor).docx"
PDF = ROOT / "assets" / "files" / "Resume (William Amor).pdf"
EMAIL = DEFAULT_EMAIL


def find_contact_paragraph(doc: Document):
    for para in doc.paragraphs[:6]:
        text = para.text
        if EMAIL in text or "linkedin.com" in text.lower():
            return para
    if len(doc.paragraphs) > 1:
        return doc.paragraphs[1]
    return None


def main() -> None:
    if not DOCX.exists():
        print(f"Missing {DOCX}", file=sys.stderr)
        sys.exit(1)

    doc = Document(str(DOCX))
    contact = find_contact_paragraph(doc)
    if contact is None:
        print("Could not find contact line paragraph.", file=sys.stderr)
        sys.exit(1)

    if len(doc.paragraphs) > 0:
        ensure_paragraph_centered(doc.paragraphs[0])
    fill_contact_line(contact, EMAIL)
    doc.save(str(DOCX))
    print(f"Updated hyperlinks in {DOCX.relative_to(ROOT)}")

    if PDF.exists():
        PDF.unlink()
    from docx2pdf import convert

    convert(str(DOCX), str(PDF))
    print(f"Wrote {PDF.relative_to(ROOT)}")


if __name__ == "__main__":
    main()

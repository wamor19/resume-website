"""
Rebuild the Word CV with clean, ATS-readable formatting.

Reads the current Word file as the source of copy, then writes it back from a
fresh document so hand-editing artefacts (mixed list styles, headings merged
into bullets, stray empty paragraphs, inherited font sizes) are gone.

Layout is single-column with no tables, text boxes, images, headers or footers,
which is what applicant tracking systems parse most reliably. Typography is
scaled down in steps until Word reports two pages; copy is never trimmed.

    python scripts/format_resume_docx.py
"""
from __future__ import annotations

import os
import re
import sys
from dataclasses import dataclass, replace
from functools import lru_cache
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parent))
from cv_hyperlinks import add_hyperlink  # noqa: E402
from word_com import (  # noqa: E402
    WD_STATISTIC_PAGES,
    WordUnavailable,
    opened,
    word_app,
)
from resume_model import (  # noqa: E402
    DOCX_PATH,
    ROOT,
    SECTION_CERTIFICATIONS,
    SECTION_EDUCATION,
    SECTION_EXPERIENCE,
    SECTION_IMPACT,
    SECTION_SUMMARY,
    SECTION_TOOLS,
    SEP,
    parse_docx,
    split_role_heading,
)

MAX_PAGES = 2
FONT = "Calibri"
_FONT_DIR = Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts"
FONT_FILES = {False: _FONT_DIR / "calibri.ttf", True: _FONT_DIR / "calibrib.ttf"}

# The contact line holds full URLs, so it is sized on its own to stay one line.
PAGE_WIDTH_PT = 8.5 * 72
CONTACT_PT_STEPS = (9.5, 9.0, 8.5, 8.0)
MIN_CONTACT_PT = CONTACT_PT_STEPS[-1]
WRAP_SAFETY_PT = 6

# Role/education headings are sized the same way: one line each, no wrapping.
META_PT_STEPS = (11.0, 10.5, 10.0, 9.5, 9.0, 8.5)
MIN_META_PT = META_PT_STEPS[-1]


@dataclass(frozen=True)
class Fit:
    """Typography knobs, tightened in steps until the CV fits two pages."""

    margin_in: float = 0.5
    name_pt: float = 19
    contact_pt: float = 9.5
    section_pt: float = 10.5
    role_title_pt: float = 10
    role_meta_pt: float = 9.5
    body_pt: float = 9.5
    line_spacing: float = 1.10
    bullet_after_pt: float = 3
    para_gap_pt: float = 6  # Between prose paragraphs, which need more air than bullets.
    para_after_pt: float = 4
    section_before_pt: float = 10
    section_after_pt: float = 4
    role_before_pt: float = 7


# Roomy first, then progressively tighter. Only spacing and point sizes change.
# Starts at Word's default Calibri 11pt and gives up readability only as needed.
FIT_STEPS: tuple[Fit, ...] = (
    Fit(
        name_pt=20,
        section_pt=12,
        role_title_pt=11.5,
        role_meta_pt=11,
        body_pt=11,
        bullet_after_pt=3.5,
        para_gap_pt=8,
        para_after_pt=5,
        section_before_pt=11,
        section_after_pt=4.5,
        role_before_pt=8,
    ),
    Fit(
        name_pt=20,
        section_pt=11.5,
        role_title_pt=11,
        role_meta_pt=10.5,
        body_pt=10.5,
        bullet_after_pt=3,
        para_gap_pt=7.5,
        para_after_pt=4.5,
        section_before_pt=10.5,
        section_after_pt=4,
        role_before_pt=7.5,
    ),
    Fit(
        section_pt=11,
        role_title_pt=10.5,
        role_meta_pt=10,
        body_pt=10,
        para_gap_pt=7,
        section_before_pt=10,
        role_before_pt=7,
    ),
    Fit(),
    Fit(line_spacing=1.06, bullet_after_pt=2.5, section_before_pt=9, role_before_pt=6),
    Fit(
        margin_in=0.45,
        line_spacing=1.04,
        bullet_after_pt=2,
        para_gap_pt=5,
        para_after_pt=3,
        section_before_pt=8,
        section_after_pt=3,
        role_before_pt=5.5,
    ),
    Fit(
        margin_in=0.45,
        section_pt=10,
        role_title_pt=9.5,
        role_meta_pt=9,
        body_pt=9,
        line_spacing=1.03,
        bullet_after_pt=2,
        para_gap_pt=4.5,
        para_after_pt=3,
        section_before_pt=7.5,
        section_after_pt=3,
        role_before_pt=5,
    ),
    Fit(
        margin_in=0.42,
        name_pt=18,
        contact_pt=9,
        section_pt=9.5,
        role_title_pt=9.5,
        role_meta_pt=8.5,
        body_pt=8.75,
        line_spacing=1.02,
        bullet_after_pt=1.5,
        para_gap_pt=4,
        para_after_pt=2.5,
        section_before_pt=7,
        section_after_pt=2.5,
        role_before_pt=4.5,
    ),
    Fit(
        margin_in=0.4,
        name_pt=17,
        contact_pt=8.5,
        section_pt=9.5,
        role_title_pt=9,
        role_meta_pt=8.5,
        body_pt=8.5,
        line_spacing=1.0,
        bullet_after_pt=1.5,
        para_gap_pt=3.5,
        para_after_pt=2,
        section_before_pt=6.5,
        section_after_pt=2,
        role_before_pt=4,
    ),
)

_URL_LIKE = re.compile(r"^(?:https?://)?[a-z0-9.-]+\.[a-z]{2,}(?:/\S*)?$", re.I)


def display_url(text: str) -> str:
    """Show URLs with their scheme so the printed text matches the link target."""
    stripped = text.rstrip("/")
    if re.match(r"^https?://", stripped, re.I):
        return stripped
    return f"https://{stripped}"


def contact_parts(headline: str) -> list[tuple[str, str | None]]:
    """Contact segments as (text, link target); target is None for plain text."""
    parts: list[tuple[str, str | None]] = []
    for part in (p.strip() for p in headline.split("|")):
        if not part:
            continue
        if "@" in part:
            parts.append((part, f"mailto:{part}"))
        elif _URL_LIKE.match(part):
            url = display_url(part)
            parts.append((url, url))
        else:
            parts.append((part, None))
    return parts


@lru_cache(maxsize=2)
def _font_metrics(bold: bool) -> tuple[dict, dict, int]:
    from fontTools.ttLib import TTFont

    font = TTFont(FONT_FILES[bold])
    return font.getBestCmap(), font["hmtx"].metrics, font["head"].unitsPerEm


def text_width_pt(text: str, size_pt: float, *, bold: bool = False) -> float | None:
    """Rendered width of text in Calibri, or None if the font can't be read."""
    try:
        cmap, metrics, upem = _font_metrics(bold)
    except Exception:  # noqa: BLE001 - font file missing or unreadable
        return None
    fallback = metrics.get("space", (upem // 2,))[0]
    units = 0
    for char in text:
        glyph = cmap.get(ord(char))
        units += metrics[glyph][0] if glyph in metrics else fallback
    return units * size_pt / upem


def usable_width_pt(fit: Fit) -> float:
    """Width of the text column, i.e. where Word starts wrapping."""
    return PAGE_WIDTH_PT - 2 * (fit.margin_in + 0.1) * 72


def fit_contact_pt(headline: str, fit: Fit) -> float:
    """Largest size keeping the contact line, with its full URLs, on one line."""
    text = SEP.join(part for part, _ in contact_parts(headline))
    usable = usable_width_pt(fit) - WRAP_SAFETY_PT
    for size in CONTACT_PT_STEPS:
        if size > fit.contact_pt:
            continue
        width = text_width_pt(text, size)
        if width is None or width <= usable:
            return size
    return min(fit.contact_pt, MIN_CONTACT_PT)


def split_education_heading(heading: str) -> tuple[str, str]:
    """"University of Kent - BSc ...' -> bold institution, plus the rest."""
    school, sep, rest = heading.partition(" - ")
    if not sep:
        school, _, rest = heading.partition(SEP)
        return school, (SEP + rest if rest else "")
    return school, (f" - {rest}" if rest else "")


def heading_lines(model: dict) -> list[tuple[str, str]]:
    """(bold part, rest) for each heading that has to stay on one line."""
    lines = []
    for role in model["roles"]:
        parts = split_role_heading(role["heading"])
        bits = [b for b in (parts["company"], *parts["extra"], parts["location"], parts["dates"]) if b]
        lines.append((parts["title"], SEP + SEP.join(bits) if bits else ""))
    if model["education"]["heading"]:
        lines.append(split_education_heading(model["education"]["heading"]))
    return lines


def heading_overflow(model: dict, fit: Fit, meta_pt: float) -> list[tuple[str, float]]:
    """Headings that would wrap at `meta_pt`, with how much too wide they are."""
    usable = usable_width_pt(fit) - WRAP_SAFETY_PT
    over_wide = []
    for bold_text, rest in heading_lines(model):
        bold_w = text_width_pt(bold_text, fit.role_title_pt, bold=True)
        rest_w = text_width_pt(rest, meta_pt)
        if bold_w is None or rest_w is None:
            return []
        over = bold_w + rest_w - usable
        if over > 0:
            over_wide.append((bold_text, over))
    return over_wide


def fit_meta_pt(model: dict, fit: Fit) -> float:
    """Largest size keeping every role/education heading on a single line."""
    for size in META_PT_STEPS:
        if size > fit.role_meta_pt:
            continue
        if not heading_overflow(model, fit, size):
            return size
    return MIN_META_PT


def _style_runs(paragraph, size_pt: float, *, bold: bool | None = None) -> None:
    for run in paragraph.runs:
        run.font.name = FONT
        run.font.size = Pt(size_pt)
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), FONT)
        if bold is not None:
            run.bold = bold


def _space(paragraph, *, before: float = 0, after: float = 0, line_spacing: float = 1.1) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line_spacing
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)


def _set_document_defaults(doc: Document, fit: Fit) -> None:
    for section in doc.sections:
        section.top_margin = Inches(fit.margin_in)
        section.bottom_margin = Inches(fit.margin_in)
        section.left_margin = Inches(fit.margin_in + 0.1)
        section.right_margin = Inches(fit.margin_in + 0.1)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(fit.body_pt)
    normal.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT)
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), "en-GB")

    # Bullets inherit Normal so ATS parsers see one consistent body font.
    bullet = doc.styles["List Bullet"]
    bullet.font.name = FONT
    bullet.font.size = Pt(fit.body_pt)


def _set_core_properties(doc: Document, model: dict) -> None:
    props = doc.core_properties
    props.author = model["name"]
    props.title = f"{model['name']} - CV"
    props.subject = "Curriculum Vitae"
    props.comments = ""
    props.keywords = ", ".join(t["items"] for t in model["tools"])[:255]


def _add_heading(doc: Document, text: str, fit: Fit) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text.upper())
    run.bold = True
    _style_runs(paragraph, fit.section_pt, bold=True)
    _space(
        paragraph,
        before=fit.section_before_pt,
        after=fit.section_after_pt,
        line_spacing=fit.line_spacing,
    )
    paragraph.paragraph_format.keep_with_next = True

    # Rule under the heading: a paragraph border, so it stays plain text for ATS.
    p_pr = paragraph._element.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "9E9E9E")
    borders.append(bottom)
    p_pr.append(borders)


# Phrases kept in italics through every rebuild (plain text in the model).
_ITALIC_PHRASES = ("vibe coding",)


def _add_run(paragraph, text: str, size_pt: float, *, bold: bool = False, italic: bool = False) -> None:
    run = paragraph.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT)


def _add_text_with_italics(paragraph, text: str, size_pt: float) -> None:
    """Write `text`, italicising known phrases such as 'vibe coding'."""
    pattern = re.compile("|".join(re.escape(p) for p in _ITALIC_PHRASES), re.I)
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            _add_run(paragraph, text[cursor : match.start()], size_pt)
        _add_run(paragraph, match.group(0), size_pt, italic=True)
        cursor = match.end()
    if cursor < len(text):
        _add_run(paragraph, text[cursor:], size_pt)


def _add_bullet(doc: Document, text: str, fit: Fit, *, last: bool) -> object:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.22)
    paragraph.paragraph_format.first_line_indent = Inches(-0.13)
    _add_text_with_italics(paragraph, text, fit.body_pt)
    _space(
        paragraph,
        after=fit.para_after_pt if last else fit.bullet_after_pt,
        line_spacing=fit.line_spacing,
    )
    paragraph.paragraph_format.widow_control = True
    return paragraph


def _write_contact_line(paragraph, headline: str, fit: Fit) -> None:
    """Rebuild the contact line, hyperlinking the email, LinkedIn and site."""
    for i, (text, target) in enumerate(contact_parts(headline)):
        if i:
            run = paragraph.add_run(SEP)
            run.font.name = FONT
            run.font.size = Pt(fit.contact_pt)
        if target:
            add_hyperlink(paragraph, text, target, size_pt=fit.contact_pt)
        else:
            run = paragraph.add_run(text)
            run.font.name = FONT
            run.font.size = Pt(fit.contact_pt)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _keep_together(paragraphs: list) -> None:
    """Stop Word splitting a role away from its first bullets."""
    for i, paragraph in enumerate(paragraphs):
        paragraph.paragraph_format.widow_control = True
        if i < len(paragraphs) - 1:
            paragraph.paragraph_format.keep_with_next = True


def build_docx(model: dict, fit: Fit) -> Document:
    doc = Document()
    _set_document_defaults(doc, fit)
    _set_core_properties(doc, model)

    name = doc.add_paragraph()
    run = name.add_run(model["name"])
    run.bold = True
    _style_runs(name, fit.name_pt, bold=True)
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _space(name, after=1.5, line_spacing=1.0)

    contact = doc.add_paragraph()
    _write_contact_line(contact, model["headline"], fit)
    _space(contact, after=fit.section_after_pt + 2, line_spacing=1.0)

    if model["summary"]:
        _add_heading(doc, SECTION_SUMMARY, fit)
        for i, para in enumerate(model["summary"]):
            paragraph = doc.add_paragraph(para)
            _style_runs(paragraph, fit.body_pt, bold=False)
            last = i == len(model["summary"]) - 1
            _space(
                paragraph,
                after=fit.para_after_pt if last else fit.para_gap_pt,
                line_spacing=fit.line_spacing,
            )

    if model["impact"]:
        _add_heading(doc, SECTION_IMPACT, fit)
        for i, item in enumerate(model["impact"]):
            _add_bullet(doc, item, fit, last=i == len(model["impact"]) - 1)

    if model["roles"]:
        _add_heading(doc, SECTION_EXPERIENCE, fit)
        for role in model["roles"]:
            block = []
            parts = split_role_heading(role["heading"])
            paragraph = doc.add_paragraph()
            title = paragraph.add_run(parts["title"])
            title.bold = True
            title.font.name = FONT
            title.font.size = Pt(fit.role_title_pt)
            meta_bits = [b for b in (parts["company"], *parts["extra"], parts["location"], parts["dates"]) if b]
            meta_text = SEP + SEP.join(meta_bits) if meta_bits else ""
            if meta_text:
                meta = paragraph.add_run(meta_text)
                meta.font.name = FONT
                meta.font.size = Pt(fit.role_meta_pt)
            _space(
                paragraph,
                before=fit.role_before_pt,
                after=2,
                line_spacing=fit.line_spacing,
            )
            block.append(paragraph)
            for i, bullet in enumerate(role["bullets"]):
                block.append(_add_bullet(doc, bullet, fit, last=i == len(role["bullets"]) - 1))
            _keep_together(block)

    education = model["education"]
    if education["heading"] or education["bullets"]:
        _add_heading(doc, SECTION_EDUCATION, fit)
        block = []
        if education["heading"]:
            # "University of Kent - BSc ... | First Class": bold the institution only.
            school, rest_text = split_education_heading(education["heading"])
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(school)
            run.bold = True
            run.font.name = FONT
            run.font.size = Pt(fit.role_title_pt)
            if rest_text:
                meta = paragraph.add_run(rest_text)
                meta.font.name = FONT
                meta.font.size = Pt(fit.role_meta_pt)
            _space(paragraph, after=2, line_spacing=fit.line_spacing)
            block.append(paragraph)
        for i, bullet in enumerate(education["bullets"]):
            block.append(_add_bullet(doc, bullet, fit, last=i == len(education["bullets"]) - 1))
        _keep_together(block)

    if model["certifications"]:
        _add_heading(doc, SECTION_CERTIFICATIONS, fit)
        certs = ", ".join(model["certifications"])
        paragraph = doc.add_paragraph(certs)
        _style_runs(paragraph, fit.body_pt, bold=False)
        _space(paragraph, after=fit.para_after_pt, line_spacing=fit.line_spacing)

    if model["tools"]:
        _add_heading(doc, SECTION_TOOLS, fit)
        block = []
        for i, group in enumerate(model["tools"]):
            paragraph = doc.add_paragraph()
            if group["label"]:
                label = paragraph.add_run(f"{group['label']}: ")
                label.bold = True
                label.font.name = FONT
                label.font.size = Pt(fit.body_pt)
            body = paragraph.add_run(group["items"])
            body.bold = False
            body.font.name = FONT
            body.font.size = Pt(fit.body_pt)
            _space(
                paragraph,
                after=fit.bullet_after_pt if i < len(model["tools"]) - 1 else 0,
                line_spacing=fit.line_spacing,
            )
            block.append(paragraph)
        _keep_together(block)

    return doc


def count_pages(path: Path) -> int | None:
    """Page count via Word automation. None when Word is unavailable."""
    try:
        with word_app() as app, opened(app, path) as doc:
            return int(doc.ComputeStatistics(WD_STATISTIC_PAGES))
    except WordUnavailable:
        return None
    except Exception:  # noqa: BLE001 - Word installed but refusing this document
        return None


def save(doc: Document, path: Path) -> None:
    try:
        doc.save(str(path))
    except PermissionError as exc:
        raise SystemExit(
            f"Cannot write {path.name} - it looks open in Word. Close it and run again."
        ) from exc


def main() -> None:
    if not DOCX_PATH.exists():
        raise SystemExit(f"Missing {DOCX_PATH}")

    model = parse_docx(DOCX_PATH)
    if not model["roles"]:
        raise SystemExit("No experience roles found in the Word file - aborting rather than overwriting it.")

    bullets = sum(len(r["bullets"]) for r in model["roles"])
    print(
        f"Parsed {len(model['roles'])} roles / {bullets} role bullets, "
        f"{len(model['impact'])} impact bullets, {len(model['summary'])} summary paragraphs."
    )

    chosen = FIT_STEPS[-1]
    for step, fit in enumerate(FIT_STEPS, start=1):
        fit = replace(
            fit,
            contact_pt=fit_contact_pt(model["headline"], fit),
            role_meta_pt=fit_meta_pt(model, fit),
        )
        chosen = fit
        save(build_docx(model, fit), DOCX_PATH)
        pages = count_pages(DOCX_PATH)
        if pages is None:
            print("Word unavailable for page count; wrote the default layout.")
            break
        print(
            f"Layout {step}/{len(FIT_STEPS)}: body {fit.body_pt}pt, headings "
            f"{fit.role_title_pt}/{fit.role_meta_pt}pt, contact {fit.contact_pt}pt -> {pages} page(s)"
        )
        if pages <= MAX_PAGES:
            break
    else:
        print(
            f"Still over {MAX_PAGES} pages at the tightest layout. Copy needs trimming.",
            file=sys.stderr,
        )
        sys.exit(1)

    for label, over in heading_overflow(model, chosen, chosen.role_meta_pt):
        print(
            f"  note: heading wraps onto a second line ({over:.0f}pt too wide): {label}",
            file=sys.stderr,
        )

    print(f"Wrote {DOCX_PATH.relative_to(ROOT)}")


if __name__ == "__main__":
    main()

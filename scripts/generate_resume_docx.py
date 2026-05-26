"""
Build Word + PDF CV from index.html (single source of truth: the site).

Run after editing resume copy on the site:
  python scripts/generate_resume_docx.py

Outputs:
  assets/files/Resume (William Amor).docx
  assets/files/Resume (William Amor).pdf  (site download link)
"""
from __future__ import annotations

import html
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

sys.path.insert(0, str(Path(__file__).resolve().parent))
from cv_fonts import (  # noqa: E402
    FONT_BODY_PT,
    FONT_CONTACT_PT,
    FONT_NAME_PT,
    FONT_ROLE_META_PT,
    FONT_ROLE_TITLE_PT,
    FONT_SECTION_PT,
)
from cv_hyperlinks import ensure_paragraph_centered, fill_contact_line  # noqa: E402

# Export trims longest roles so Word/PDF fit two pages; the site keeps full bullets.
DEFAULT_BULLET_CAPS: dict[str, int] = {
    "Principal Product Owner": 7,
    "Data & Analytics Product Owner (EMEA)": 6,
    "IT Business Technology Leader (Leadership Programme)": 5,
    "Intelligent Automation Product Manager (Leadership Programme)": 7,
    "Global Project Coordinator": 4,
    "Technology Business Partner (Industrial Placement & Part-time)": 3,
}
EXPORT_BULLET_CAPS: dict[str, int] = {}
PAGE2_ROLE_INDEX = 3  # Johnson & Johnson Automation — everything from here starts page 2

ROOT = Path(__file__).resolve().parents[1]
HTML_PATH = ROOT / "index.html"
DOCX_PATH = ROOT / "assets" / "files" / "Resume (William Amor).docx"
PDF_PATH = ROOT / "assets" / "files" / "Resume (William Amor).pdf"


_ASCII_FOLD = {
    "\u2018": "'",   # left single quote
    "\u2019": "'",   # right single quote / apostrophe
    "\u201A": ",",   # single low-9 quote
    "\u201C": '"',   # left double quote
    "\u201D": '"',   # right double quote
    "\u201E": '"',   # double low-9 quote
    "\u2013": "-",   # en dash
    "\u2014": "-",   # em dash
    "\u2026": "...",  # ellipsis
    "\u2011": "-",   # non-breaking hyphen
    "\u00a0": " ",   # nbsp
    "\u2022": "-",   # bullet
    "\u00b7": "|",   # middle dot (used as separator on site)
    "\u00b7 ": "| ",
}


def ascii_fold(text: str) -> str:
    for src, dst in _ASCII_FOLD.items():
        text = text.replace(src, dst)
    return text


def strip_html(text: str) -> str:
    text = re.sub(r"<br\s*/?>", " ", text, flags=re.I)
    text = re.sub(r"<[^>]+>", "", text)
    text = html.unescape(text)
    text = ascii_fold(text)
    return re.sub(r"\s+", " ", text).strip()


def first_match(pattern: str, source: str, flags=0) -> str:
    m = re.search(pattern, source, flags)
    return strip_html(m.group(1)) if m else ""


def first_match_html(pattern: str, source: str, flags=0) -> str:
    m = re.search(pattern, source, flags)
    return m.group(1) if m else ""


def all_matches(pattern: str, source: str, flags=0) -> list[str]:
    return [strip_html(m.group(1)) for m in re.finditer(pattern, source, flags)]


def parse_resume_html(source: str) -> dict:
    email = first_match(r'data-email="([^"]+)"', source)

    ledes = all_matches(r'<p class="hero__lede">\s*(.*?)\s*</p>', source, re.DOTALL)
    summary = " ".join(ledes) if ledes else ""

    impact_block = first_match_html(
        r'<ul class="impactWins"[^>]*>(.*?)</ul>', source, re.DOTALL
    )
    impact = all_matches(r"<li>(.*?)</li>", impact_block, re.DOTALL) if impact_block else []

    roles = []
    for card in re.finditer(
        r'<article class="xpCard"[^>]*data-tags="([^"]*)"[^>]*>(.*?)</article>',
        source,
        re.DOTALL,
    ):
        block = card.group(2)
        title = first_match(r'<h3 class="xpCard__title">(.*?)</h3>', block, re.DOTALL)
        company = first_match(
            r'<span class="xpCard__orgName">(.*?)</span>', block, re.DOTALL
        )
        location = first_match(
            r'<span class="xpCard__orgLoc">(.*?)</span>', block, re.DOTALL
        )
        dates = first_match(r'<div class="xpCard__date">(.*?)</div>', block, re.DOTALL)
        dates = dates.replace(" to present", " - Present").replace(" to ", " - ")
        bullets_html = first_match_html(
            r'<ul class="bullets">\s*(.*?)\s*</ul>', block, re.DOTALL
        )
        bullet_items = (
            all_matches(r"<li>(.*?)</li>", bullets_html, re.DOTALL) if bullets_html else []
        )
        roles.append(
            {
                "title": title,
                "company": company,
                "location": location,
                "dates": dates,
                "bullets": bullet_items,
            }
        )

    edu_section = first_match_html(
        r'<section class="section" id="education"[^>]*>(.*?)</section>',
        source,
        re.DOTALL,
    )
    grade_html = first_match_html(
        r'<div class="eduChip__grade">(.*?)</div>', edu_section, re.DOTALL
    ) or ""
    # Drop the visual separator span (middle dot etc.) for ATS-clean output
    grade_html = re.sub(
        r'\s*<span class="eduChip__gradeSep"[^>]*>.*?</span>\s*',
        ", ",
        grade_html,
        flags=re.DOTALL,
    )
    education = {
        "school": first_match(r'<div class="eduChip__h">(.*?)</div>', edu_section, re.DOTALL),
        "degree": first_match(
            r'<div class="eduChip__degree">(.*?)</div>', edu_section, re.DOTALL
        ),
        "grade": strip_html(grade_html),
        "bullets": [],
    }
    edu_bullets_html = first_match_html(
        r'<ul class="eduChip__bullets"[^>]*>(.*?)</ul>', edu_section, re.DOTALL
    )
    education["bullets"] = (
        all_matches(r"<li>(.*?)</li>", edu_bullets_html, re.DOTALL) if edu_bullets_html else []
    )

    cert_block = first_match_html(
        r'<section class="section" id="certificates"[^>]*>.*?<ul class="bullets"[^>]*>(.*?)</ul>',
        source,
        re.DOTALL,
    )
    certificates = (
        all_matches(r"<li>(.*?)</li>", cert_block, re.DOTALL) if cert_block else []
    )

    skills_section = first_match_html(
        r'<section class="section" id="toolbox-skills"[^>]*>(.*?)</section>',
        source,
        re.DOTALL,
    )
    skill_groups: list[tuple[str, str]] = []
    if skills_section:
        tool_group_pat = (
            r'<h3 class="toolGroup__h"[^>]*>(.*?)</h3>\s*'
            r'<div class="tags"[^>]*>(.*?)</div>'
        )
        legacy_flow_pat = (
            r'<h3 class="skillFlow__h"[^>]*>(.*?)</h3>\s*'
            r'<p class="skillFlow__p">\s*(.*?)\s*</p>'
        )
        for pattern in (tool_group_pat, legacy_flow_pat):
            for m in re.finditer(pattern, skills_section, re.DOTALL):
                label = strip_html(m.group(1))
                body_block = m.group(2)
                if '<span class="tag">' in body_block:
                    tags = all_matches(
                        r'<span class="tag">(.*?)</span>', body_block, re.DOTALL
                    )
                    body = ", ".join(tags)
                else:
                    body = strip_html(body_block)
                skill_groups.append((label, body))
            if skill_groups:
                break

    chips = all_matches(r'<span class="chip">(.*?)</span>', source, re.DOTALL)

    return {
        "email": email,
        "summary": summary,
        "impact": impact,
        "roles": roles,
        "education": education,
        "certificates": certificates,
        "skill_groups": skill_groups,
        "chips": chips,
    }


def set_doc_defaults(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.45)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)


def style_paragraph(
    paragraph,
    size_pt: float,
    *,
    space_before: float = 0,
    space_after: float = 3,
    line_spacing: float = 1.12,
) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line_spacing
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        run.font.size = Pt(size_pt)


def bullets_for_export(title: str, bullets: list[str]) -> list[str]:
    cap = EXPORT_BULLET_CAPS.get(title)
    if cap is None or len(bullets) <= cap:
        return bullets
    return bullets[:cap]


def add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(FONT_SECTION_PT)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True


def keep_block_together(paragraphs: list) -> None:
    """Keep a block (role, education, etc.) on one page when Word can fit it."""
    for i, para in enumerate(paragraphs):
        para.paragraph_format.keep_together = True
        para.paragraph_format.widow_control = True
        if i < len(paragraphs) - 1:
            para.paragraph_format.keep_with_next = True


def add_role(doc: Document, role: dict, *, page_break_before: bool = False) -> None:
    block: list = []

    p = doc.add_paragraph()
    if page_break_before:
        p.paragraph_format.page_break_before = True
    title_run = p.add_run(role["title"])
    title_run.bold = True
    title_run.font.size = Pt(FONT_ROLE_TITLE_PT)
    title_run.font.name = "Calibri"
    meta = p.add_run(
        f"  |  {role['company']}  |  {role['location']}  |  {role['dates']}"
    )
    meta.font.size = Pt(FONT_ROLE_META_PT)
    meta.font.name = "Calibri"
    style_paragraph(p, FONT_ROLE_TITLE_PT, space_before=7, space_after=2)
    block.append(p)

    export_bullets = bullets_for_export(role["title"], role["bullets"])
    for j, bullet in enumerate(export_bullets):
        bp = doc.add_paragraph(bullet, style="List Bullet")
        bp.paragraph_format.left_indent = Inches(0.2)
        bp.paragraph_format.first_line_indent = Inches(-0.12)
        style_paragraph(
            bp,
            FONT_BODY_PT,
            space_after=3 if j < len(export_bullets) - 1 else 5,
        )
        block.append(bp)

    keep_block_together(block)


def build_docx(data: dict) -> Document:
    doc = Document()
    set_doc_defaults(doc)

    h = doc.add_paragraph()
    name = h.add_run("William Amor")
    name.bold = True
    name.font.size = Pt(FONT_NAME_PT)
    name.font.name = "Calibri"
    style_paragraph(h, FONT_NAME_PT, space_after=2)
    ensure_paragraph_centered(h)

    sub = doc.add_paragraph()
    fill_contact_line(sub, data["email"] or "message@william-amor.info", size_pt=FONT_CONTACT_PT)
    style_paragraph(sub, FONT_CONTACT_PT, space_after=10)

    add_section_heading(doc, "Professional Summary")
    sp = doc.add_paragraph(data["summary"])
    style_paragraph(sp, FONT_CONTACT_PT, space_after=4, line_spacing=1.15)

    if data["impact"]:
        add_section_heading(doc, "Impact Highlights")
        for i, item in enumerate(data["impact"]):
            bp = doc.add_paragraph(item, style="List Bullet")
            bp.paragraph_format.left_indent = Inches(0.2)
            bp.paragraph_format.first_line_indent = Inches(-0.12)
            style_paragraph(
                bp,
                FONT_BODY_PT,
                space_after=3 if i < len(data["impact"]) - 1 else 5,
            )

    add_section_heading(doc, "Experience")
    for i, role in enumerate(data["roles"]):
        add_role(doc, role, page_break_before=(i == PAGE2_ROLE_INDEX))

    edu = data["education"]
    if edu.get("school"):
        add_section_heading(doc, "Education")
        edu_block: list = []
        p = doc.add_paragraph()
        school = p.add_run(edu["school"])
        school.bold = True
        school.font.size = Pt(FONT_ROLE_TITLE_PT)
        school.font.name = "Calibri"
        degree_text = edu.get("degree", "")
        grade_text = edu.get("grade", "")
        suffix_parts = [t for t in (degree_text, grade_text) if t]
        if suffix_parts:
            line = p.add_run("  -  " + "  |  ".join(suffix_parts))
            line.font.size = Pt(FONT_ROLE_META_PT)
            line.font.name = "Calibri"
        edu_block.append(p)
        style_paragraph(p, FONT_CONTACT_PT, space_after=3)
        edu_bullets = edu.get("bullets", [])
        for j, bullet in enumerate(edu_bullets):
            bp = doc.add_paragraph(bullet, style="List Bullet")
            bp.paragraph_format.left_indent = Inches(0.2)
            bp.paragraph_format.first_line_indent = Inches(-0.12)
            style_paragraph(
                bp,
                FONT_BODY_PT,
                space_after=3 if j < len(edu_bullets) - 1 else 4,
            )
            edu_block.append(bp)
        keep_block_together(edu_block)

    if data["certificates"]:
        add_section_heading(doc, "Certifications")
        cert_block: list = []
        for j, cert in enumerate(data["certificates"]):
            bp = doc.add_paragraph(cert, style="List Bullet")
            bp.paragraph_format.left_indent = Inches(0.2)
            bp.paragraph_format.first_line_indent = Inches(-0.12)
            style_paragraph(
                bp,
                FONT_BODY_PT,
                space_after=2 if j < len(data["certificates"]) - 1 else 4,
            )
            cert_block.append(bp)
        keep_block_together(cert_block)

    if data["skill_groups"]:
        add_section_heading(doc, "Tools & Platforms")
        skills_block: list = []
        for j, (label, body) in enumerate(data["skill_groups"]):
            p = doc.add_paragraph()
            head = p.add_run(f"{label}: ")
            head.bold = True
            head.font.size = Pt(FONT_BODY_PT)
            head.font.name = "Calibri"
            rest = p.add_run(body)
            rest.font.size = Pt(FONT_BODY_PT)
            rest.font.name = "Calibri"
            style_paragraph(
                p,
                FONT_BODY_PT,
                space_after=3 if j < len(data["skill_groups"]) - 1 else 2,
                line_spacing=1.15,
            )
            skills_block.append(p)
        keep_block_together(skills_block)

    return doc


def count_docx_pages(docx_path: Path) -> int | None:
    """Return page count using Word (Windows). None if Word automation unavailable."""
    try:
        import win32com.client
    except ImportError:
        return None

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc = None
    try:
        doc = word.Documents.Open(str(docx_path.resolve()))
        return int(doc.ComputeStatistics(2))
    finally:
        if doc is not None:
            doc.Close(False)
        word.Quit()


def reset_export_caps() -> None:
    EXPORT_BULLET_CAPS.clear()
    EXPORT_BULLET_CAPS.update(DEFAULT_BULLET_CAPS)


def tighten_export_caps() -> None:
    for key in list(EXPORT_BULLET_CAPS.keys()):
        EXPORT_BULLET_CAPS[key] = max(3, EXPORT_BULLET_CAPS[key] - 1)


def export_pdf(docx_path: Path, pdf_path: Path) -> None:
    try:
        from docx2pdf import convert
    except ImportError as exc:
        raise SystemExit("docx2pdf is required: pip install docx2pdf") from exc

    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    if pdf_path.exists():
        pdf_path.unlink()
    convert(str(docx_path), str(pdf_path))


def main() -> None:
    if not HTML_PATH.exists():
        print(f"Missing {HTML_PATH}", file=sys.stderr)
        sys.exit(1)

    source = HTML_PATH.read_text(encoding="utf-8")
    data = parse_resume_html(source)
    if not data["roles"]:
        print("No experience roles parsed from index.html", file=sys.stderr)
        sys.exit(1)

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    reset_export_caps()

    for attempt in range(5):
        doc = build_docx(data)
        doc.save(DOCX_PATH)
        pages = count_docx_pages(DOCX_PATH)
        if pages is not None:
            print(f"Word pages: {pages}")
            if pages <= 2:
                break
            tighten_export_caps()
            print("Over 2 pages — tightening export bullet caps and rebuilding...", file=sys.stderr)
        else:
            break
    else:
        print("Could not fit resume to 2 Word pages after retries.", file=sys.stderr)
        sys.exit(1)

    print(f"Wrote {DOCX_PATH.relative_to(ROOT)}")

    from cv_fonts import apply_fonts_to_document  # noqa: E402

    doc = Document(str(DOCX_PATH))
    if len(doc.paragraphs) > 1:
        fill_contact_line(
            doc.paragraphs[1],
            data["email"] or "message@william-amor.info",
            size_pt=FONT_CONTACT_PT,
        )
    apply_fonts_to_document(doc)
    doc.save(str(DOCX_PATH))

    try:
        export_pdf(DOCX_PATH, PDF_PATH)
        print(f"Wrote {PDF_PATH.relative_to(ROOT)}")
    except Exception as exc:  # noqa: BLE001
        print(f"PDF export failed ({exc}). Open Word and export PDF manually.", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()

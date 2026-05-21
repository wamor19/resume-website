"""Insert Education section into the Word CV (before Certifications), then export PDF."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

sys.path.insert(0, str(Path(__file__).resolve().parent))
from cv_fonts import (  # noqa: E402
    FONT_BODY_PT,
    FONT_CONTACT_PT,
    FONT_ROLE_META_PT,
    FONT_ROLE_TITLE_PT,
    FONT_SECTION_PT,
    apply_fonts_to_document,
)
from generate_resume_docx import (  # noqa: E402
    DOCX_PATH,
    PDF_PATH,
    count_docx_pages,
    export_pdf,
    keep_block_together,
    style_paragraph,
)
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]

EDUCATION = {
    "school": "University of Kent",
    "degree": "BSc (Hons) Computing and Business Administration with a year in industry",
    "grade": "First Class Honours  |  US GPA equivalent: ~4.0",
    "bullets": [
        "Year in Industry at GSK, then extended into final year: 20 hrs/week part-time on an RPA product automating goods ordering in a manufacturing plant.",
        'Dissertation (First Class): "The disruption and emergence of FinTech".',
    ],
}


def insert_paragraph_before(ref: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    ref._p.addprevious(new_p)
    return Paragraph(new_p, ref._parent)


def insert_section_heading_before(ref: Paragraph, text: str) -> Paragraph:
    p = insert_paragraph_before(ref)
    run = p.add_run(text.upper())
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(FONT_SECTION_PT)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    return p


def find_certifications_paragraph(doc: Document) -> Paragraph | None:
    for para in doc.paragraphs:
        if para.text.strip().upper() == "CERTIFICATIONS":
            return para
    return None


def has_education_section(doc: Document) -> bool:
    return any(p.text.strip().upper() == "EDUCATION" for p in doc.paragraphs)


def add_education(doc: Document, anchor: Paragraph) -> None:
    block: list[Paragraph] = []
    cursor = anchor

    for bullet in reversed(EDUCATION["bullets"]):
        bp = insert_paragraph_before(cursor)
        bp.style = doc.styles["List Bullet"]
        bp.add_run(bullet)
        bp.paragraph_format.left_indent = Inches(0.2)
        bp.paragraph_format.first_line_indent = Inches(-0.12)
        style_paragraph(bp, FONT_BODY_PT, space_after=3)
        block.insert(0, bp)
        cursor = bp

    p = insert_paragraph_before(cursor)
    school = p.add_run(f"{EDUCATION['school']} — ")
    school.bold = True
    school.font.size = Pt(FONT_ROLE_TITLE_PT)
    school.font.name = "Calibri"
    line = p.add_run(f"{EDUCATION['degree']}  |  {EDUCATION['grade']}")
    line.font.size = Pt(FONT_ROLE_META_PT)
    line.font.name = "Calibri"
    style_paragraph(p, FONT_CONTACT_PT, space_after=3)
    block.insert(0, p)
    cursor = p

    heading = insert_section_heading_before(cursor, "Education")
    block.insert(0, heading)
    keep_block_together(block)


def main() -> None:
    if not DOCX_PATH.exists():
        raise SystemExit(f"Missing {DOCX_PATH}")

    doc = Document(str(DOCX_PATH))
    if has_education_section(doc):
        print("Education section already present — skipping insert.")
    else:
        anchor = find_certifications_paragraph(doc)
        if anchor is None:
            raise SystemExit("Could not find CERTIFICATIONS heading in docx.")
        add_education(doc, anchor)
        print(f"Added Education to {DOCX_PATH.relative_to(ROOT)}")

    apply_fonts_to_document(doc)
    doc.save(str(DOCX_PATH))

    pages = count_docx_pages(DOCX_PATH)
    if pages is not None:
        print(f"Word pages: {pages}")

    export_pdf(DOCX_PATH, PDF_PATH)
    print(f"Wrote {PDF_PATH.relative_to(ROOT)}")


if __name__ == "__main__":
    main()

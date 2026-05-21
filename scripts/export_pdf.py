"""Export the current Word CV to the site PDF (Word is source of truth)."""
import sys
from pathlib import Path

from docx import Document
from docx2pdf import convert

sys.path.insert(0, str(Path(__file__).resolve().parent))
from cv_fonts import apply_fonts_to_document  # noqa: E402
from cv_hyperlinks import DEFAULT_EMAIL, ensure_paragraph_centered, fill_contact_line  # noqa: E402
from add_cv_hyperlinks import find_contact_paragraph  # noqa: E402

EMAIL = DEFAULT_EMAIL

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "assets" / "files" / "Resume (William Amor).docx"
PDF = ROOT / "assets" / "files" / "Resume (William Amor).pdf"

def ensure_hyperlinks() -> None:
    doc = Document(str(DOCX))
    contact = find_contact_paragraph(doc)
    if len(doc.paragraphs) > 0:
        ensure_paragraph_centered(doc.paragraphs[0])
    if contact is not None:
        fill_contact_line(contact, EMAIL)
    apply_fonts_to_document(doc)
    doc.save(str(DOCX))


if __name__ == "__main__":
    if not DOCX.exists():
        raise SystemExit(f"Missing {DOCX}")
    ensure_hyperlinks()
    if PDF.exists():
        PDF.unlink()
    convert(str(DOCX), str(PDF))
    print(f"Wrote {PDF.relative_to(ROOT)}")
